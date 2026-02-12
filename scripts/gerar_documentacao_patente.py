#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Script para gerar documentação técnica do Sistema Matriz CSV to KML
para fins de registro de patente ou proteção de propriedade intelectual.

Gera um arquivo Word (.docx) com descrição completa do sistema, funcionalidades e funções.
"""
import os
import sys

# Adiciona o diretório raiz ao path
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

def criar_documentacao():
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
    except ImportError:
        print("ERRO: Instale python-docx: pip install python-docx")
        sys.exit(1)

    doc = Document()
    
    # Configuração de estilos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    def titulo(texto, nivel=1):
        p = doc.add_paragraph()
        run = p.add_run(texto)
        if nivel == 1:
            run.bold = True
            run.font.size = Pt(16)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        elif nivel == 2:
            run.bold = True
            run.font.size = Pt(13)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
        elif nivel == 3:
            run.bold = True
            run.font.size = Pt(11)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
        return p

    def paragrafo(texto):
        return doc.add_paragraph(texto)

    # ========== CAPA E IDENTIFICAÇÃO ==========
    titulo("DOCUMENTAÇÃO TÉCNICA PARA REGISTRO DE PATENTE", 1)
    paragrafo("Sistema Matriz CSV to KML – Sistema Integrado de Projeto e Dimensionamento de Redes de Distribuição Elétrica")
    doc.add_paragraph()
    paragrafo("Documento confidencial – Propriedade intelectual")
    doc.add_paragraph()
    paragrafo("Este documento descreve de forma completa as funcionalidades, componentes, algoritmos e métodos implementados no sistema, com fins de registro e proteção de patente.")

    # ========== RESUMO PARA PEDIDO DE PATENTE (INPI) ==========
    titulo("RESUMO PARA PEDIDO DE PATENTE", 1)
    paragrafo(
        "A presente invenção refere-se a um sistema informatizado para o projeto e dimensionamento automatizado de redes de distribuição de energia elétrica de média e baixa tensão. "
        "O sistema compreende: (a) módulo de entrada de dados configurado para receber traçados de linhas de transmissão sob a forma de polilinhas geográficas, seja por desenho manual em mapa interativo, importação de arquivos KML, GeoJSON ou matriz CSV; "
        "(b) módulo de processamento geodésico que calcula distâncias, ângulos de deflexão e posiciona postes intercalados conforme parâmetros de vão médio e tramo máximo; "
        "(c) módulo de ábacos técnicos com polígonos de decisão que determinam automaticamente o tipo de estrutura, poste e equipamentos a partir do ângulo de deflexão e das distâncias entre vértices consecutivos; "
        "(d) módulo de encabeçamento automático que identifica pontos de mudança de direção que exigem estruturas reforçadas; e "
        "(e) módulo de exportação que gera arquivos CSV com coordenadas geográficas e UTM, KML para visualização em SIG e DXF para uso em software CAD. "
        "O sistema permite definir segmentos onde não se intercalam postes e suporta o vão frouxo na origem do traçado. "
        "A solução reduz erros, acelera o dimensionamento e integra de forma inédita mapa, matriz, ábacos e exportação para CAD/SIG em fluxo único.")
    paragrafo("Aplicação industrial: concessionárias de distribuição, empresas de projetos de redes e escritórios de engenharia que realizam levantamentos topográficos e dimensionamento de redes de distribuição de energia elétrica.")
    paragrafo("Palavras-chave: distribuição elétrica; rede de média tensão; projeto automatizado; ábaco técnico; KML; DXF; dimensionamento de postes; encabeçamento; polilinha geográfica; CAD; SIG.")

    # ========== 1. RESUMO DO SISTEMA ==========
    titulo("1. RESUMO DO SISTEMA", 1)
    paragrafo(
        "O Sistema Matriz CSV to KML é uma solução informatizada para o projeto e dimensionamento automatizado de redes de distribuição de energia elétrica. "
        "O sistema recebe traçados de linhas de transmissão ou polilinhas (via arquivos KML, GeoJSON ou entrada manual no mapa) e gera automaticamente: "
        "(i) matrizes de postes com coordenadas geográficas e UTM; "
        "(ii) posicionamento de postes intercalados considerando vão médio e tramo máximo; "
        "(iii) atribuição de estruturas (cabos, postes, bases) mediante consulta a ábacos técnicos; "
        "(iv) identificação automática de encabeçamentos por ângulo de deflexão; "
        "(v) exportação em formatos CSV, KML e DXF para uso em CAD e SIG.")
    paragrafo(
        "O sistema inclui interface web com mapa interativo (Leaflet), API REST (Flask), e processamento backend em Python, "
        "sendo utilizável tanto a partir de desenho de polilinha no mapa quanto a partir de importação de arquivos ou matriz CSV existente.")

    # ========== 2. OBJETIVO E CAMPO DE APLICAÇÃO ==========
    titulo("2. OBJETIVO E CAMPO DE APLICAÇÃO", 1)
    paragrafo(
        "O sistema tem como objetivo automatizar o processo de projeto de redes de distribuição elétrica de média e baixa tensão, "
        "reduzindo erros humanos e tempo de dimensionamento. Aplica-se a:")
    paragrafo("• Concessionárias de energia elétrica")
    paragrafo("• Empresas de projetos de redes")
    paragrafo("• Escritórios de engenharia e topografia")
    paragrafo("• Integração com sistemas CAD e SIG")

    # ========== 3. ARQUITETURA DO SISTEMA ==========
    titulo("3. ARQUITETURA DO SISTEMA", 1)
    titulo("3.1. Componentes Principais", 2)
    paragrafo("O sistema é composto por três camadas:")
    paragrafo("• Frontend: Interface web com mapa interativo (Leaflet), formulários e tabela de dados de módulos (ábacos)")
    paragrafo("• Backend: Processamento em Python (Flask) com cálculos geográficos, ábacos e exportação")
    paragrafo("• Armazenamento: Arquivos CSV, KML, DXF e Excel para entrada/saída")

    titulo("3.2. Fluxos de Uso", 2)
    paragrafo("Fluxo 1 – A partir de polilinha no mapa: Usuário desenha a linha no mapa ou importa KML/HTML; seleciona módulo, parâmetros e segmentos onde não intercalar; o sistema gera CSV, KML e DXF com postes dimensionados.")
    paragrafo("Fluxo 2 – A partir de CSV importado: Usuário importa matriz CSV existente; clica em Plotar Projeto; o sistema processa o CSV, gera KML e DXF e exibe o traçado no mapa.")

    # ========== 4. FUNCIONALIDADES DETALHADAS ==========
    titulo("4. FUNCIONALIDADES DETALHADAS", 1)

    titulo("4.1. Entrada de Dados", 2)
    paragrafo("• Importação de KML/KMZ: Extração automática de vértices de LineString, Polygon e Point")
    paragrafo("• Importação de GeoJSON/HTML: Suporte a múltiplos formatos de coordenadas")
    paragrafo("• Desenho manual no mapa: Cliques para definir vértices, inversão de sentido, finalização da polilinha")
    paragrafo("• Importação de matriz CSV: Aceita CSV com múltiplas linhas por vértice ou uma linha por vértice; transformação automática")
    paragrafo("• Tabela de módulos: Seleção de módulo (ábaco) com vão médio, tramo máximo e demais parâmetros")
    paragrafo("• Segmentos sem intercalação: Definição de vértices iniciais onde não devem ser intercalados postes")

    titulo("4.2. Cálculos Geográficos", 2)
    paragrafo("• Distância Haversine: Cálculo de distância em metros entre dois pontos em WGS84")
    paragrafo("• Ângulo (bearing): Direção entre dois pontos em graus (0–360)")
    paragrafo("• Ponto polar: Novo ponto a partir de distância e ângulo")
    paragrafo("• Ângulo de deflexão: Ângulo entre três pontos consecutivos (0–180°) para decisão de encabeçamento")
    paragrafo("• Conversão UTM: Conversão WGS84 para coordenadas UTM (fuso, easting, northing) para CAD")

    titulo("4.3. Processamento de Vértices", 2)
    paragrafo("• Vão frouxo (Loose Gap): Inserção de ponto intermediário a 30 m quando distância entre os dois primeiros pontos > 60 m")
    paragrafo("• Divisão de tramo: Inserção de pontos intermediários quando o tramo excede o vão máximo (section_size)")
    paragrafo("• Marcação por ângulo de deflexão: Consulta ao ábaco (mosaico) para definir se o vértice é encabeçamento")
    paragrafo("• Encabeçamento automático: Marcação de encabeçamento quando distância acumulada excede section_size")
    paragrafo("• Intercalação de postes: Inserção de postes intermediários quando distância > gap_size, exceto em segmentos configurados")

    titulo("4.4. Ábacos e Estruturas", 2)
    paragrafo("• Ábacos (mtz_abaco): Dicionários de estruturas por código (10101, 10102, 10104, 10105, etc.) com polígonos de decisão")
    paragrafo("• Mosaico (mosaico, point_in_polygon): Consulta de estrutura por distância1, distância2 e ângulo de deflexão usando polígonos")
    paragrafo("• Atribuição de estruturas: tipo_poste, estru_mt_nv1, base_concreto, estai_ancora, rotacao_poste, etc.")
    paragrafo("• Suporte a múltiplos módulos com mesmo ábaco (codigo_abaco)")

    titulo("4.5. Exportação", 2)
    paragrafo("• CSV: Matriz com sequência, status, lat, long, estruturas, fuso, utm_x, utm_y e colunas adicionais; separador ; e decimal ,")
    paragrafo("• KML: Quadrados 5×3 m na bissetriz, linhas conectando vértices, estilos por status (implantar, existente, retirar, deslocar)")
    paragrafo("• DXF R12 (AC1009): Pontos e polylines em UTM, textos com informações de poste, compatível com CAD")
    paragrafo("• Coordenadas UTM no CSV: Colunas fuso, utm_x, utm_y nas últimas colunas para cada ponto implantar")

    # ========== 5. FUNÇÕES E MÓDULOS ==========
    titulo("5. CATÁLOGO DE FUNÇÕES E MÓDULOS", 1)

    titulo("5.1. Backend – Core (Cálculos e Processamento)", 2)
    paragrafo("• gerar_matriz(): Função principal que orquestra todo o processamento; recebe trecho, módulo, parâmetros e vértices; retorna DataFrame da matriz")
    paragrafo("• distance(), angle(), polar(), distance_ptos(), angulo_deflexao(): Funções de cálculo geográfico (calculo_geografico.py)")
    paragrafo("• get_loose_gap(): Aplica vão frouxo quando habilitado (processamento_vertices.py)")
    paragrafo("• dividir_tramo(): Divide tramos longos em seções menores (processamento_vertices.py)")
    paragrafo("• intercalar_vertices(): Insere postes intermediários (processamento_vertices.py)")
    paragrafo("• transformar_csv_para_uma_linha_por_vertice(): Converte CSV múltiplas linhas/vértice para uma linha/vértice (transformacao_csv.py)")

    titulo("5.2. Backend – Elementos", 2)
    paragrafo("• marcar_vertices_angulo_deflexao(): Marca encabeçamentos por ângulo e distância via ábaco")
    paragrafo("• colocar_encabecamento_rede(): Encabeçamento automático por distância acumulada")
    paragrafo("• colocar_poste_estrutura(): Atribui estruturas e postes a cada vértice usando mosaico")
    paragrafo("• gravar_pontos_matriz(): Registra dados de estruturas no dicionário de pontos")

    titulo("5.3. Backend – Ábacos", 2)
    paragrafo("• mtz_abaco(): Retorna lista de entradas do ábaco por código")
    paragrafo("• mosaico(): Consulta estrutura por distância1, distância2 e ângulo")
    paragrafo("• point_in_polygon(): Verifica se ponto está dentro de polígono de decisão")

    titulo("5.4. Backend – Exportação", 2)
    paragrafo("• criar_kml_quadrados_bissetriz(): Gera KML com quadrados na bissetriz e linhas")
    paragrafo("• criar_dxf_do_kml(): Gera DXF R12 a partir do KML")
    paragrafo("• latlon_to_utm_with_zone(): Converte lat/lon para (fuso, easting, northing)")
    paragrafo("• salvar_csv(): Persiste DataFrame em CSV no disco")
    paragrafo("• exportar_para_kml(): Exporta pontos para KML básico")

    titulo("5.5. API REST (Flask)", 2)
    paragrafo("• POST /api/gerar-matriz/: Recebe vértices e parâmetros; retorna CSV, KML e DXF em base64")
    paragrafo("• POST /api/plotar-projeto-csv/: Recebe CSV importado; retorna KML e DXF em base64")
    paragrafo("• GET /api/test/: Teste de disponibilidade da API")

    titulo("5.6. Frontend (JavaScript)", 2)
    paragrafo("• initMap(): Inicialização do mapa Leaflet")
    paragrafo("• loadKMLOnMap(), loadGeoJSONOnMap(): Carregamento de KML/GeoJSON no mapa")
    paragrafo("• Modo manual: Criar polilinha por cliques, btnFinalizarPolilinha, btnInverterSentido")
    paragrafo("• Gerar Matriz: Chamada à API com vértices do mapa; download de CSV, KML, DXF")
    paragrafo("• Importar Matriz CSV / Plotar Projeto: Upload de CSV; plotagem e download de KML/DXF")
    paragrafo("• showGeneratedKmlOnMap(): Exibe KML gerado no mapa com popups")
    paragrafo("• buildGeoJsonFromKml(): Converte KML para GeoJSON com suporte a namespace")

    # ========== 6. INOVAÇÕES E DIFERENCIAIS ==========
    titulo("6. INOVAÇÕES E DIFERENCIAIS TÉCNICOS", 1)
    paragrafo("• Integração completa mapa ↔ matriz ↔ KML ↔ DXF em um único fluxo")
    paragrafo("• Ábacos parametrizados com polígonos de decisão (mosaico) para seleção automática de estrutura")
    paragrafo("• Encabeçamento automático por ângulo de deflexão e distância acumulada")
    paragrafo("• Suporte a segmentos sem intercalação configuráveis por vértice")
    paragrafo("• Conversão automática para UTM no CSV para uso em CAD")
    paragrafo("• Transformação de CSV múltiplas linhas/vértice para formato normalizado")
    paragrafo("• Suporte a namespace KML para compatibilidade com diferentes geradores")

    # ========== 7. TECNOLOGIAS UTILIZADAS ==========
    titulo("7. TECNOLOGIAS UTILIZADAS", 1)
    paragrafo("• Linguagens: Python 3, JavaScript")
    paragrafo("• Backend: Flask, Pandas, lxml, openpyxl")
    paragrafo("• Frontend: HTML5, CSS, Leaflet, PapaParse, SheetJS (xlsx)")
    paragrafo("• Formatos: CSV (sep ;), KML 2.2, DXF R12 (AC1009), GeoJSON")
    paragrafo("• Coordenadas: WGS84, projeção UTM (fuso automático)")

    # ========== 8. GLOSSÁRIO ==========
    titulo("8. GLOSSÁRIO", 1)
    paragrafo("• Vão frouxo: Ponto intermediário inserido no início do traçado quando a distância inicial é grande")
    paragrafo("• Vão médio / section_size: Distância máxima da seção entre encabeçamentos")
    paragrafo("• Tramo máximo / gap_size: Distância máxima entre postes consecutivos")
    paragrafo("• Encabeçamento: Poste que suporta mudança de direção (ângulo de deflexão)")
    paragrafo("• Ábaco: Tabela de decisão que associa ângulos e distâncias a estruturas específicas")
    paragrafo("• Mosaico: Conjunto de polígonos no plano (dist1, dist2) para escolha de estrutura")

    # Salvar documento
    base_dir = os.path.dirname(os.path.dirname(__file__))
    output_path = os.path.join(base_dir, "Documentacao_Patente_Sistema_Matriz_CSV_to_KML.docx")
    try:
        doc.save(output_path)
    except PermissionError:
        output_path = os.path.join(base_dir, "Documentacao_Patente_Sistema_Matriz_CSV_to_KML_v2.docx")
        doc.save(output_path)
    print(f"Documento gerado: {output_path}")
    return output_path


if __name__ == "__main__":
    criar_documentacao()
